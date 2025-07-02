#!/usr/bin/env python3
import json
import uuid
import re
import urllib.request
import urllib.parse
import io
import tempfile
import os
from html import unescape
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs

# 尝试导入文档解析库
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not found. Word document parsing will be limited.")

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    print("Warning: PyPDF2 not found. PDF parsing will be limited.")

try:
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False
    print("Warning: openpyxl not found. Excel parsing will be limited.")

# ---------------- 全局内存存储 ----------------
# 在简易示例中，用于持久化会话级数据，避免 list 接口返回空数组。
# 生产环境应替换为数据库或持久层。

# 数据集结构示例
# {
#   'id': 'uuid',
#   'name': 'kb_xxxxxxxx',
#   'description': str,
#   'created_at': ISO8601 字符串,
#   'documents': [ { 'id': 'docId', 'name': 'document_xx', 'status': 'completed' } ]
# }

# ---------------- 持久化 ----------------

DATA_FILE = '/app/data/datasets.json'

def _load_datasets():
    try:
        if os.path.exists(DATA_FILE):
            with open(DATA_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        print(f'Failed to load datasets: {e}')
    return {}

def _save_datasets():
    try:
        os.makedirs(os.path.dirname(DATA_FILE), exist_ok=True)
        with open(DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(DATASETS, f, ensure_ascii=False)
    except Exception as e:
        print(f'Failed to save datasets: {e}')

# 初始化
DATASETS = _load_datasets()

class SimpleRAGHandler(BaseHTTPRequestHandler):
    def extract_text_from_docx(self, file_data):
        """从Word文档中提取文本内容"""
        if not HAS_DOCX:
            return "Word文档解析需要安装python-docx库: pip install python-docx"
        
        try:
            # 使用内存中的数据创建Document对象
            doc = Document(io.BytesIO(file_data))
            
            # 提取标题
            title = "Word文档"
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                title = doc.paragraphs[0].text.strip()[:100]  # 使用第一段作为标题
            
            # 提取所有段落文本
            content_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content.append(cell_text)
                    if row_content:
                        table_content.append(" | ".join(row_content))
                if table_content:
                    content_parts.append("\n".join(table_content))
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                content = "无法提取到文档内容"
            
            return {
                'title': title,
                'markdown': content
            }
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            return {
                'title': "Word文档解析错误",
                'markdown': f"解析Word文档时出错: {str(e)}"
            }
    
    def extract_text_from_pdf(self, file_data):
        """从PDF文档中提取文本内容"""
        if not HAS_PDF:
            return "PDF解析需要安装PyPDF2库: pip install PyPDF2"
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
            
            title = f"PDF文档 ({len(pdf_reader.pages)}页)"
            content_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        content_parts.append(f"--- 第{page_num + 1}页 ---\n{page_text.strip()}")
                except Exception as e:
                    content_parts.append(f"--- 第{page_num + 1}页 ---\n无法提取页面内容: {str(e)}")
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                content = "无法提取到PDF内容"
            
            return {
                'title': title,
                'markdown': content
            }
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            return {
                'title': "PDF解析错误",
                'markdown': f"解析PDF时出错: {str(e)}"
            }
    
    def extract_text_from_excel(self, file_data):
        """从Excel文档中提取文本内容"""
        if not HAS_EXCEL:
            return "Excel解析需要安装openpyxl库: pip install openpyxl"
        
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_data))
            
            title = f"Excel文档 ({len(workbook.sheetnames)}个工作表)"
            content_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content_parts.append(f"--- 工作表: {sheet_name} ---")
                
                # 获取有数据的区域
                rows_data = []
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空行
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):
                        rows_data.append(" | ".join(row_data))
                
                if rows_data:
                    content_parts.append("\n".join(rows_data))
                else:
                    content_parts.append("(无数据)")
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                content = "无法提取到Excel内容"
            
            return {
                'title': title,
                'markdown': content
            }
        except Exception as e:
            print(f"Error parsing Excel: {e}")
            return {
                'title': "Excel解析错误",
                'markdown': f"解析Excel时出错: {str(e)}"
            }
    
    def scrape_url(self, url):
        """增强的文档解析功能，支持多种文件格式"""
        try:
            print(f"Processing URL: {url}")
            
            # 创建请求
            req = urllib.request.Request(url)
            req.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
            
            # 发送请求并获取数据
            with urllib.request.urlopen(req, timeout=30) as response:
                content_type = response.headers.get('Content-Type', '').lower()
                file_data = response.read()
            
            print(f"Content-Type: {content_type}, Data size: {len(file_data)} bytes")
            
            # 根据Content-Type或URL扩展名判断文件类型
            url_lower = url.lower()
            
            # Word文档处理
            if ('application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type or 
                url_lower.endswith('.docx')):
                print("Detected Word document (.docx)")
                return self.extract_text_from_docx(file_data)
            
            # PDF文档处理  
            elif ('application/pdf' in content_type or url_lower.endswith('.pdf')):
                print("Detected PDF document")
                return self.extract_text_from_pdf(file_data)
            
            # Excel文档处理
            elif ('application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' in content_type or
                  url_lower.endswith('.xlsx')):
                print("Detected Excel document (.xlsx)")
                return self.extract_text_from_excel(file_data)
            
            # 纯文本文件
            elif ('text/' in content_type or url_lower.endswith(('.txt', '.md', '.csv'))):
                print("Detected text document")
                try:
                    content = file_data.decode('utf-8', errors='ignore')
                    title = "文本文档"
                    
                    # 使用第一行作为标题（如果存在）
                    lines = content.split('\n', 1)
                    if lines and lines[0].strip():
                        title = lines[0].strip()[:100]
                    
                    return {
                        'title': title,
                        'markdown': content.strip()
                    }
                except Exception as e:
                    return {
                        'title': "文本解析错误",
                        'markdown': f"解析文本文件时出错: {str(e)}"
                    }
            
            # HTML/网页内容处理（改进的逻辑）
            else:
                print("Processing as HTML/web content", flush=True)
                try:
                    raw_content = file_data.decode('utf-8', errors='ignore')
                    
                    # 提取标题
                    title = "网页文档"
                    
                    # 首先尝试从<title>标签提取
                    title_match = re.search(r'<title[^>]*>(.*?)</title>', raw_content, re.IGNORECASE | re.DOTALL)
                    if title_match:
                        title = unescape(title_match.group(1).strip())
                        # 清理标题中的换行和多余空格
                        title = re.sub(r'\s+', ' ', title).strip()
                    else:
                        # 如果没有title标签，尝试从h1标签提取
                        h1_match = re.search(r'<h1[^>]*>(.*?)</h1>', raw_content, re.IGNORECASE | re.DOTALL)
                        if h1_match:
                            h1_text = re.sub(r'<[^>]+>', '', h1_match.group(1))
                            title = unescape(h1_text.strip())[:100]
                    
                    # 移除脚本和样式
                    content = re.sub(r'<script[^>]*>.*?</script>', '', raw_content, flags=re.IGNORECASE | re.DOTALL)
                    content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.IGNORECASE | re.DOTALL)
                    content = re.sub(r'<noscript[^>]*>.*?</noscript>', '', content, flags=re.IGNORECASE | re.DOTALL)
                    
                    # 处理常见HTML元素，转换为更好的文本格式
                    content = re.sub(r'<br[^>]*>', '\n', content, flags=re.IGNORECASE)
                    content = re.sub(r'<p[^>]*>', '\n\n', content, flags=re.IGNORECASE)
                    content = re.sub(r'</p>', '', content, flags=re.IGNORECASE)
                    content = re.sub(r'<h[1-6][^>]*>', '\n\n# ', content, flags=re.IGNORECASE)
                    content = re.sub(r'</h[1-6]>', '\n', content, flags=re.IGNORECASE)
                    content = re.sub(r'<li[^>]*>', '\n- ', content, flags=re.IGNORECASE)
                    content = re.sub(r'</li>', '', content, flags=re.IGNORECASE)
                    
                    # 移除剩余的HTML标签
                    content = re.sub(r'<[^>]+>', '', content)
                    
                    # 清理HTML实体
                    content = unescape(content)
                    
                    # 清理文本格式
                    content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)  # 合并多个空行
                    content = re.sub(r'[ \t]+', ' ', content)  # 合并多个空格
                    content = content.strip()
                    
                    # 限制内容长度
                    if len(content) > 15000:
                        content = content[:15000] + "\n\n[内容已截断...]"
                    
                    # 如果提取的内容为空，使用原始内容的片段
                    if len(content.strip()) < 10:
                        content = raw_content[:1000] + "..." if len(raw_content) > 1000 else raw_content
                        title = "内容解析失败"
                    
                    return {
                        'title': title,
                        'markdown': content
                    }
                except Exception as e:
                    return {
                        'title': "内容解析错误",
                        'markdown': f"解析内容时出错: {str(e)}"
                    }
                    
        except Exception as e:
            print(f"Error processing URL {url}: {e}")
            return {
                'title': f"处理失败: {os.path.basename(url)}",
                'markdown': f"无法处理文档 {url}。错误: {str(e)}"
            }

    def do_POST(self):
        # 解析请求路径
        path = urlparse(self.path).path
        
        if path == '/api/v1/scrape':
            # 处理爬虫请求
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length > 0:
                post_data = self.rfile.read(content_length)
                try:
                    data = json.loads(post_data.decode('utf-8'))
                    url = data.get('url', '')
                    kb_id = data.get('kb_id', '')
                    
                    print(f"Scraping URL: {url} for KB: {kb_id}", flush=True)
                    
                    # 执行爬取和解析
                    scraped_data = self.scrape_url(url)
                    
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json')
                    self.send_header('Access-Control-Allow-Origin', '*')
                    self.end_headers()
                    
                    response = {
                        'err': 0,
                        'msg': 'success',
                        'data': scraped_data
                    }
                    self.wfile.write(json.dumps(response).encode())
                    print(f"Successfully processed: {scraped_data['title']}")
                    return
                except Exception as e:
                    print(f"Error processing scrape request: {e}")
                    self.send_response(500)
                    self.send_header('Content-type', 'application/json')
                    self.send_header('Access-Control-Allow-Origin', '*')
                    self.end_headers()
                    response = {
                        'err': 1,
                        'msg': f'Failed to process request: {str(e)}',
                        'data': {}
                    }
                    self.wfile.write(json.dumps(response).encode())
                    return
        
        elif path == '/api/v1/datasets':
            # 创建知识库
            # 读取并丢弃上传内容，避免客户端阻塞
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length > 0:
                _ = self.rfile.read(content_length)

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            dataset_id = str(uuid.uuid4())
            dataset_obj = {
                'id': dataset_id,
                'name': 'kb_' + dataset_id[:8],
                'description': 'Auto-created knowledge base',
                'created_at': '2025-06-25T08:00:00Z'
            }

            # 存入内存
            dataset_copy = dataset_obj.copy()
            dataset_copy['documents'] = []
            DATASETS[dataset_id] = dataset_copy
            _save_datasets()

            response = {
                'code': 0,
                'data': dataset_obj
            }

            self.wfile.write(json.dumps(response).encode('utf-8'))
            print(f"Created dataset: {dataset_obj}")
            return
        
        elif path.startswith('/api/v1/datasets/') and '/documents' in path:
            # 创建文档
            # 路径格式: /api/v1/datasets/{id}/documents or /api/v1/datasets/{id}/documents/text
            parts = path.strip('/').split('/')
            dataset_id = parts[3] if len(parts) > 3 else ''
            if dataset_id not in DATASETS:
                self.send_response(404)
                self.send_header('Content-type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(json.dumps({'err': 1, 'msg': 'dataset not found'}).encode())
                return
            # 读取 multipart 内容，粗略解析文件文本
            content_length = int(self.headers.get('Content-Length', 0))
            body_bytes = self.rfile.read(content_length) if content_length > 0 else b''
            file_text = ''
            try:
                import re, base64
                ctype = self.headers.get('Content-Type', '')
                boundary_match = re.search('boundary=(.*)', ctype)
                if boundary_match:
                    boundary = boundary_match.group(1).encode()
                    parts_bin = body_bytes.split(b'--' + boundary)
                    for part in parts_bin:
                        if b'Content-Disposition' in part and b'\r\n\r\n' in part:
                            file_content = part.split(b'\r\n\r\n', 1)[1]
                            # 去掉结尾的换行及boundary尾
                            file_content = file_content.rstrip(b'\r\n')
                            # 尝试解码
                            try:
                                file_text = file_content.decode('utf-8', errors='ignore')
                            except Exception:
                                file_text = base64.b64encode(file_content).decode()
                            break
            except Exception as e:
                print(f'Failed to parse multipart upload: {e}')
            # fallback
            if not file_text:
                file_text = 'dummy uploaded content'

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            doc_id = str(uuid.uuid4())
            doc_obj = {
                'id': doc_id,
                'dataset_id': dataset_id,
                'name': 'document_' + doc_id[:8],
                'status': 'completed',
                'content': file_text,
            }
            DATASETS[dataset_id]['documents'].append(doc_obj)
            _save_datasets()

            response = {
                'code': 0,
                'data': [doc_obj]
            }
            self.wfile.write(json.dumps(response).encode('utf-8'))
            print(f"Created document: {doc_obj['id']} (len={len(file_text)} chars)")
            return
        
        elif path == '/api/v1/model-configs':
            # 添加模型配置
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length > 0:
                post_data = self.rfile.read(content_length)
                try:
                    data = json.loads(post_data.decode('utf-8'))
                    print(f"Model config request: {data}")
                except:
                    pass
            
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            model_id = str(uuid.uuid4())
            response = {
                'id': model_id,
                'name': 'model_' + model_id[:8],
                'status': 'active'
            }
            self.wfile.write(json.dumps(response).encode())
            print(f"Added model config: {response}")
            return
        
        elif path.startswith('/api/v1/datasets/') and path.endswith('/chunks'):
            # 解析文档（将全文作为单块）
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length) if content_length > 0 else b'{}'
            try:
                payload = json.loads(body.decode('utf-8'))
            except:
                payload = {}
            parts = path.strip('/').split('/')
            dataset_id = parts[3]
            doc_ids = payload.get('document_ids', [])
            for doc_id in doc_ids:
                # 查找对应文档
                doc_list = [d for d in DATASETS[dataset_id].get('documents', []) if d['id'] == doc_id]
                if not doc_list:
                    continue
                content_text = doc_list[0].get('content', '')
                chunk = {
                    'id': str(uuid.uuid4()),
                    'content': content_text[:2048],  # 限制长度
                    'document_id': doc_id,
                    'dataset_id': dataset_id,
                    'similarity': 1.0,
                }
                DATASETS[dataset_id].setdefault('chunks', []).append(chunk)
            _save_datasets()

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            response = {'code': 0}
            self.wfile.write(json.dumps(response).encode('utf-8'))
            print(f"Parsed documents for dataset {dataset_id}: {doc_ids}")
            return
        
        elif path == '/api/v1/retrieval':
            # 简易检索：根据question关键字匹配chunks
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length) if content_length > 0 else b'{}'
            req = json.loads(body.decode('utf-8'))
            dataset_ids = req.get('dataset_ids', [])
            question = req.get('question', '').lower()
            result_chunks = []
            for ds_id in dataset_ids:
                ds = DATASETS.get(ds_id)
                if not ds:
                    continue
                for ch in ds.get('chunks', []):
                    # 关键词简单匹配
                    if not question or question in ch.get('content', '').lower():
                        result_chunks.append(ch)
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            resp = {
                'code': 0,
                'data': {
                    'chunks': result_chunks,
                    'total': len(result_chunks)
                }
            }
            self.wfile.write(json.dumps(resp).encode('utf-8'))
            print(f"Retrieval request for datasets {dataset_ids}, question='{question}' -> {len(result_chunks)} chunks")
            return
        
        # 默认返回成功响应
        self.send_response(200)
        self.send_header('Content-type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        response = {'success': True, 'message': 'Operation completed'}
        self.wfile.write(json.dumps(response).encode())
        print(f"POST {path} - returned default success")
    
    def do_PUT(self):
        # 处理PUT请求（更新操作）
        path = urlparse(self.path).path
        
        self.send_response(200)
        self.send_header('Content-type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        response = {'success': True, 'message': 'Update completed'}
        self.wfile.write(json.dumps(response).encode())
        print(f"PUT {path} - returned success")
    
    def do_DELETE(self):
        # 处理DELETE请求（删除操作）
        path = urlparse(self.path).path
        
        self.send_response(200)
        self.send_header('Content-type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        response = {'success': True, 'message': 'Delete completed'}
        self.wfile.write(json.dumps(response).encode())
        print(f"DELETE {path} - returned success")
    
    def do_GET(self):
        path = urlparse(self.path).path
        
        if path == '/api/v1/datasets':
            # 数据集列表，支持 name/id 过滤
            query = parse_qs(urlparse(self.path).query)
            name_filter = query.get('name', [''])[0]
            id_filter = query.get('id', [''])[0]

            datasets = []
            for ds in DATASETS.values():
                if name_filter and name_filter != ds['name']:
                    continue
                if id_filter and id_filter != ds['id']:
                    continue
                datasets.append(ds)

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            resp = {'code': 0, 'data': datasets}
            self.wfile.write(json.dumps(resp).encode('utf-8'))
            print(f"List datasets: count={len(datasets)}")
            return

        if path.startswith('/api/v1/datasets/'):
            parts = path.strip('/').split('/')
            if len(parts) == 3:
                dataset_id = parts[2]
                ds = DATASETS.get(dataset_id)
                if ds is None:
                    self.send_response(404)
                    self.end_headers()
                    return
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(json.dumps({'code':0,'data': ds}).encode('utf-8'))
                return

        if path == '/api/v1/models':
            # 返回模型配置列表 - 这是后端真正访问的端点
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            # 后端期望ListModelConfigsResponse格式: {code: int, data: []ModelConfig}
            response = {
                'code': 0,
                'data': [
                    {
                        'id': 'embedding-config-1',
                        'provider': 'openai-compatible-api',
                        'name': 'Default Embedding Model',
                        'task_type': 'embedding',
                        'api_base': 'https://api.siliconflow.cn/v1',
                        'api_key': 'mock-key',
                        'max_tokens': 8192,
                        'is_default': True,
                        'enabled': True,
                        'config': {},
                        'description': 'Default embedding model configuration',
                        'version': '1.0',
                        'timeout': 30,
                        'create_time': 1735114800,
                        'update_time': 1735114800,
                        'owner': 'system',
                        'quota_limit': 1000
                    },
                    {
                        'id': 'rerank-config-1',
                        'provider': 'openai-compatible-api',
                        'name': 'Default Rerank Model',
                        'task_type': 'rerank',
                        'api_base': 'https://api.siliconflow.cn/v1',
                        'api_key': 'mock-key',
                        'max_tokens': 4096,
                        'is_default': True,
                        'enabled': True,
                        'config': {},
                        'description': 'Default rerank model configuration',
                        'version': '1.0',
                        'timeout': 30,
                        'create_time': 1735114800,
                        'update_time': 1735114800,
                        'owner': 'system',
                        'quota_limit': 1000
                    }
                ]
            }
            self.wfile.write(json.dumps(response).encode())
            print("Returned model config list via /models endpoint")
            return
        
        elif path == '/api/v1/model-configs':
            # 返回模型配置列表 - 符合ListModelConfigsResponse格式
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            # 后端期望ListModelConfigsResponse格式: {code: int, data: []ModelConfig}
            response = {
                'code': 0,
                'data': [
                    {
                        'id': 'embedding-config-1',
                        'provider': 'openai-compatible-api',
                        'name': 'Default Embedding Model',
                        'task_type': 'embedding',
                        'api_base': 'https://api.siliconflow.cn/v1',
                        'api_key': 'mock-key',
                        'max_tokens': 8192,
                        'is_default': True,
                        'enabled': True,
                        'config': {},
                        'description': 'Default embedding model configuration',
                        'version': '1.0',
                        'timeout': 30,
                        'create_time': 1735114800,
                        'update_time': 1735114800,
                        'owner': 'system',
                        'quota_limit': 1000
                    },
                    {
                        'id': 'rerank-config-1',
                        'provider': 'openai-compatible-api',
                        'name': 'Default Rerank Model',
                        'task_type': 'rerank',
                        'api_base': 'https://api.siliconflow.cn/v1',
                        'api_key': 'mock-key',
                        'max_tokens': 4096,
                        'is_default': True,
                        'enabled': True,
                        'config': {},
                        'description': 'Default rerank model configuration',
                        'version': '1.0',
                        'timeout': 30,
                        'create_time': 1735114800,
                        'update_time': 1735114800,
                        'owner': 'system',
                        'quota_limit': 1000
                    }
                ]
            }
            self.wfile.write(json.dumps(response).encode())
            print("Returned model config list (ListModelConfigsResponse format)")
            return
        
        elif path == '/api/v1/datasets':
            # 返回数据集列表
            # 解析查询参数
            query = urlparse(self.path).query
            params = parse_qs(query)
            filter_name = params.get('name', [None])[0]
            page_size = int(params.get('page_size', [100])[0])

            all_ds = list(DATASETS.values())
            if filter_name:
                all_ds = [d for d in all_ds if d['name'] == filter_name]

            # 按创建时间排序并截断
            result = [
                {
                    'id': d['id'],
                    'name': d['name'],
                    'description': d['description'],
                    'created_at': d['created_at']
                } for d in all_ds[:page_size]
            ]

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps(result).encode())
            print("Returned dataset list, count:", len(result))
            return

        elif path.startswith('/api/v1/datasets/'):
            # 处理按 ID 获取数据集或其文档
            parts = path.strip('/').split('/')
            if len(parts) == 3:  # /api/v1/datasets/{id}
                dataset_id = parts[2]
                ds = DATASETS.get(dataset_id)
                if ds:
                    resp = {
                        'id': ds['id'],
                        'name': ds['name'],
                        'description': ds['description'],
                        'created_at': ds['created_at']
                    }
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json')
                    self.send_header('Access-Control-Allow-Origin', '*')
                    self.end_headers()
                    self.wfile.write(json.dumps(resp).encode())
                else:
                    self.send_response(404)
                    self.send_header('Access-Control-Allow-Origin', '*')
                    self.end_headers()
                return
            elif '/documents' in path:
                dataset_id = parts[3] if len(parts) > 3 else ''
                ds = DATASETS.get(dataset_id)
                if ds:
                    self.send_response(200)
                    self.send_header('Content-type', 'application/json')
                    self.send_header('Access-Control-Allow-Origin', '*')
                    self.end_headers()
                    self.wfile.write(json.dumps(ds['documents']).encode())
                else:
                    self.send_response(404)
                    self.send_header('Access-Control-Allow-Origin', '*')
                    self.end_headers()
                return
        
        elif path == '/health' or path == '/healthz':
            # 健康检查
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            response = {'status': 'healthy'}
            self.wfile.write(json.dumps(response).encode())
            return
        
        # 默认返回404
        self.send_response(404)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        print(f"GET {path} - returned 404")
    
    def do_OPTIONS(self):
        # 处理CORS预检请求
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, PUT, DELETE, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        self.end_headers()
        
    def log_message(self, format, *args):
        # 自定义日志格式
        print(f"[RAG] {format % args}")

if __name__ == '__main__':
    print('Starting enhanced RAG service on port 8080...')
    print('Supported endpoints:')
    print('  GET  /api/v1/models')
    print('  GET  /api/v1/model-configs')
    print('  POST /api/v1/model-configs')
    print('  GET  /api/v1/datasets')
    print('  POST /api/v1/datasets')
    print('  POST /api/v1/datasets/{id}/documents')
    print('  GET  /health')
    
    server = HTTPServer(('0.0.0.0', 8080), SimpleRAGHandler)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print('Shutting down...')
        server.shutdown() 